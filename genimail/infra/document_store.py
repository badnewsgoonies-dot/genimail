"""File-system operations for quote documents.

This module owns all I/O that was previously in ``genimail.domain.quotes``.
The domain layer stays pure (no ``os`` / ``shutil`` / ``webbrowser``).
"""

import os
import shutil
import webbrowser
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from genimail.domain.quotes import _sanitize_filename_part, render_quote_template_text
from genimail.paths import DEFAULT_QUOTE_TEMPLATE_FILE, QUOTE_DIR

_DOC_EXTENSIONS = {".doc", ".docx"}


def _create_docx_template(template_path: str):
    """Generate the default GX Painting quote template as a formatted .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Company header + date on same line ---
    p = doc.add_paragraph()
    run = p.add_run("GX Painting LTD")
    run.font.size = Pt(11)
    run.font.color.rgb = None
    p.add_run("\t\t\t\t\t\t\t[Date Here]")

    for line in [
        "25 Robert Berry Cres",
        "King City ON",
        "L7B0M4",
        "Cell: (416)-560-8741",
        "Email: gxpainting@hotmail.com",
    ]:
        p = doc.add_paragraph(line)

    doc.add_paragraph()  # blank line

    # --- Quotation heading ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quotation")
    run.bold = True

    doc.add_paragraph()  # blank line

    # --- Submitted to / Job Name ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted to:")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Job Name:")
    run.bold = True

    doc.add_paragraph()  # blank line

    # --- Below find the estimate ---
    p = doc.add_paragraph("Below find the estimate for.")
    p.runs[0].bold = True

    doc.add_paragraph()  # blank line

    # --- Scope of Work (bold + underline) ---
    p = doc.add_paragraph()
    run = p.add_run("Scope of Work")
    run.bold = True
    run.underline = True

    doc.add_paragraph()  # blank line

    # --- Bullet point ---
    doc.add_paragraph("\t", style="List Bullet")

    doc.add_paragraph()  # blank line

    # --- Total Price ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Total Price $ () + HST")
    run.bold = True

    doc.add_paragraph()  # blank line

    # --- Separate Price / Note ---
    p = doc.add_paragraph()
    run = p.add_run("Separate Price:")
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Note:")
    run.bold = True

    doc.add_paragraph()  # blank line

    # --- Disclaimer ---
    doc.add_paragraph(
        "We supply all labour, materials, supervisions, tools, equipment "
        "and insurance necessary to carry out and complete all of the "
        "painting at the above project."
    )

    doc.add_paragraph()  # blank line

    # --- Signature ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Gaz Xure")

    doc.save(template_path)


def ensure_default_quote_template(template_path: str):
    """Create the default quote template if one does not exist."""
    if os.path.exists(template_path):
        return
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    if template_path.lower().endswith(".docx"):
        _create_docx_template(template_path)
    else:
        # Fallback for non-.docx paths (legacy plain-text with placeholders)
        with open(template_path, "w", encoding="utf-8", newline="\r\n") as f:
            f.write(
                "GX Painting LTD\n"
                "Quote #: {{QUOTE_ID}}\n"
                "Date: {{DATE}}\n"
                "\nClient: {{CLIENT_NAME}}\n"
                "Email: {{CLIENT_EMAIL}}\n"
                "Project: {{PROJECT_NAME}}\n"
                "Reference: {{EMAIL_SUBJECT}}\n"
                "\nPrepared by: GX Painting LTD\n"
            )


def create_quote_doc(template_path: str, output_dir: str, context: dict) -> str:
    """Read *template_path*, render placeholders from *context*, write to *output_dir*."""
    template_path = os.path.abspath((template_path or "").strip() or DEFAULT_QUOTE_TEMPLATE_FILE)
    output_dir = os.path.abspath((output_dir or "").strip() or QUOTE_DIR)
    ensure_default_quote_template(template_path)
    os.makedirs(output_dir, exist_ok=True)

    client_name = context.get("client_name") or context.get("client_email") or "Client"
    client_slug = _sanitize_filename_part(client_name, fallback="Client")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    with open(template_path, "rb") as f:
        template_bytes = f.read()

    template_ext = os.path.splitext(template_path)[1].lower()
    if template_bytes.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        out_ext = ".doc"
    elif template_bytes.startswith(b"PK\x03\x04"):
        if template_ext in {".docx", ".docm", ".dotx", ".dotm"}:
            out_ext = template_ext
        else:
            out_ext = ".docx"
    else:
        out_ext = ".doc"

    out_path = os.path.join(output_dir, f"Quote_{stamp}_{client_slug}{out_ext}")

    if template_bytes.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        shutil.copy2(template_path, out_path)
        return out_path
    if template_bytes.startswith(b"PK\x03\x04"):
        shutil.copy2(template_path, out_path)
        return out_path

    template_text = None
    for enc in ("utf-8-sig", "utf-16"):
        try:
            template_text = template_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue

    if template_text is None:
        try:
            cp_text = template_bytes.decode("cp1252")
            printable = sum(ch.isprintable() or ch in "\r\n\t" for ch in cp_text)
            if cp_text and printable / len(cp_text) > 0.9:
                template_text = cp_text
        except UnicodeDecodeError:
            pass

    if template_text is None:
        shutil.copy2(template_path, out_path)
        return out_path

    rendered = render_quote_template_text(template_text, context)
    with open(out_path, "w", encoding="utf-8", newline="\r\n") as f:
        f.write(rendered)
    return out_path


def open_document_file(path: str) -> bool:
    """Open a document with the OS default handler."""
    target = os.path.abspath(path)
    try:
        os.startfile(target)
        return True
    except Exception:
        try:
            webbrowser.open(f"file:///{target.replace(os.sep, '/')}")
            return True
        except Exception:
            return False


def latest_doc_file(output_dir: str):
    """Return the most-recently-modified ``.doc``/``.docx`` file in *output_dir*, or ``None``."""
    if not output_dir or not os.path.isdir(output_dir):
        return None
    matches = []
    for name in os.listdir(output_dir):
        if os.path.splitext(name)[1].lower() in _DOC_EXTENSIONS:
            full = os.path.join(output_dir, name)
            try:
                matches.append((os.path.getmtime(full), full))
            except OSError:
                pass
    if not matches:
        return None
    matches.sort(key=lambda item: item[0], reverse=True)
    return matches[0][1]
